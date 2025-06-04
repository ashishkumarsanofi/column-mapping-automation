"""
Convert the architecture guide from Markdown to Word document format.
This script creates a professional Word document with proper formatting.
"""

import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    print("python-docx library is available")
except ImportError:
    print("python-docx library not found. Installing...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

def create_word_document():
    """Create a professional Word document from the architecture guide."""
    
    # Create a new document
    doc = Document()
    
    # Set up document properties
    doc.core_properties.title = "Advanced Column Mapping & Transformation Tool - Architecture Guide"
    doc.core_properties.author = "Technical Documentation Team"
    doc.core_properties.subject = "Software Architecture Documentation"
    doc.core_properties.keywords = "Python, Streamlit, Data Processing, Architecture"
    
    # Add title page
    title = doc.add_heading('📊 Advanced Column Mapping & Transformation Tool', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Complete Architecture & Flow Guide', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add document info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run("Document Version: 1.0\n").bold = True
    info_para.add_run("Date: June 4, 2025\n")
    info_para.add_run("Author: Technical Documentation\n")
    
    # Add page break
    doc.add_page_break()
    
    # Table of Contents
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_para = doc.add_paragraph()
    toc_items = [
        "1. Executive Summary",
        "2. File Structure & Responsibilities", 
        "3. Application Flow Sequence",
        "4. Inter-Module Dependencies",
        "5. Key Processing Points",
        "6. User Journey Flow",
        "7. Technical Details",
        "8. Enhanced Features"
    ]
    for item in toc_items:
        toc_para.add_run(item + "\n")
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    exec_summary = doc.add_paragraph(
        "The Advanced Column Mapping & Transformation Tool is a Streamlit-based web application "
        "designed to help users map, transform, and consolidate data from multiple input files "
        "into a single output file. The tool provides an intuitive interface for managing column "
        "mappings, applying filters, handling static values, and formatting data with enhanced "
        "mapping file support for configuration reusability."
    )
    
    # Key Capabilities
    doc.add_heading('Key Capabilities:', level=3)
    capabilities = [
        "Multi-file data consolidation",
        "Dynamic column mapping", 
        "Advanced filtering capabilities",
        "Static value assignment",
        "Date formatting controls",
        "Enhanced mapping file format with backward compatibility",
        "Real-time error validation",
        "Multiple export formats (Excel, TXT, CSV)"
    ]
    for cap in capabilities:
        p = doc.add_paragraph(cap, style='List Bullet')
    
    # File Structure & Responsibilities
    doc.add_heading('2. File Structure & Responsibilities', level=1)
    
    # app.py section
    doc.add_heading('2.1 app.py - Main Orchestrator (Entry Point)', level=2)
    doc.add_paragraph("Primary Role: Central controller that coordinates the entire application workflow", style='Intense Quote')
    
    doc.add_heading('Key Responsibilities:', level=3)
    app_responsibilities = [
        "Application Configuration: Sets up Streamlit page configuration, layout, and styling",
        "File Upload Management: Handles file uploads for input files, output templates, and mapping configurations",
        "Sheet Selection Logic: Manages Excel sheet selection for multi-sheet workbooks", 
        "Mapping File Validation: Validates and processes both basic and enhanced mapping file formats",
        "Workflow Coordination: Orchestrates the complete data processing pipeline",
        "Enhanced Features Support: Detects and enables advanced mapping capabilities"
    ]
    for resp in app_responsibilities:
        doc.add_paragraph(resp, style='List Bullet')
    
    # ui_sections.py section
    doc.add_heading('2.2 ui_sections.py - User Interface Components', level=2)
    doc.add_paragraph("Primary Role: Provides reusable UI components for clean code separation", style='Intense Quote')
    
    doc.add_heading('Key Responsibilities:', level=3)
    ui_responsibilities = [
        "File Upload Interface: Creates the file upload section with proper styling and validation",
        "User Guide Display: Renders comprehensive help documentation with examples",
        "Footer Management: Displays application footer and branding",
        "UI Consistency: Ensures consistent styling and user experience across components"
    ]
    for resp in ui_responsibilities:
        doc.add_paragraph(resp, style='List Bullet')
    
    # file_utils.py section
    doc.add_heading('2.3 file_utils.py - File Operations & Utilities', level=2)
    doc.add_paragraph("Primary Role: Handles all file reading and data validation operations", style='Intense Quote')
    
    doc.add_heading('Key Responsibilities:', level=3)
    file_responsibilities = [
        "File Reading: Efficiently reads CSV and Excel files with optimized performance",
        "Data Validation: Handles file format validation and error reporting",
        "Column Management: Ensures required columns exist in DataFrames",
        "Error Handling: Provides comprehensive error reporting for file-related issues"
    ]
    for resp in file_responsibilities:
        doc.add_paragraph(resp, style='List Bullet')
    
    # mapping_logic.py section
    doc.add_heading('2.4 mapping_logic.py - Core Business Logic Engine', level=2)
    doc.add_paragraph("Primary Role: The heart of the application - contains all mapping and processing logic", style='Intense Quote')
    
    doc.add_heading('Key Responsibilities:', level=3)
    mapping_responsibilities = [
        "Mapping Interface Creation: Builds dynamic tabbed interface for each input file/sheet",
        "Enhanced Mapping Support: Processes and applies enhanced mapping configurations",
        "Data Processing: Handles filtering, static value assignment, and date formatting",
        "Output Generation: Creates final consolidated output with multiple export options",
        "Configuration Export: Exports current mapping configurations for reuse",
        "Real-time Validation: Provides immediate feedback on mapping errors"
    ]
    for resp in mapping_responsibilities:
        doc.add_paragraph(resp, style='List Bullet')
    
    # Application Flow Sequence
    doc.add_heading('3. Application Flow Sequence', level=1)
    
    doc.add_heading('Phase 1: Initialization & Setup', level=2)
    phase1_text = """
    1. Application Startup (app.py)
       • Import all required dependencies (streamlit, pandas, custom modules)
       • Configure Streamlit application (page title, layout, colors)
       • Set maximum upload size limits
       • Display application header and introduction
    """
    doc.add_paragraph(phase1_text)
    
    doc.add_heading('Phase 2: File Upload & Validation', level=2)
    phase2_text = """
    2. File Upload Process (ui_sections.py → app.py)
       • Display upload interface via show_upload_section()
       • Accept input files (multiple CSV/Excel files)
       • Accept output template (single CSV/Excel file)
       • Accept optional mapping file (CSV/Excel with mapping config)
       • Validate uploaded files and report errors
    
    3. Sheet Selection Logic (app.py)
       • Process Excel files to extract sheet information
       • Display multiselect widget for each Excel file
       • Allow users to select specific sheets for processing
       • Handle CSV files (no sheet selection needed)
       • Build input_file_sheets list with all combinations
    """
    doc.add_paragraph(phase2_text)
    
    doc.add_heading('Phase 3: Mapping File Processing', level=2)
    phase3_text = """
    4. Mapping File Analysis (app.py)
       • Read and validate mapping file structure
       • Check for required basic columns (FileName, SheetName, OutputColumn, InputColumn)
       • Detect enhanced features (StaticValue, FilterValues, DateFormatFlag, IncludeFlag)
       • Add missing columns with defaults for backward compatibility
       • Notify user when enhanced features are detected
       • Set mapping_file_valid flag for downstream processing
    """
    doc.add_paragraph(phase3_text)
    
    doc.add_heading('Phase 4: Core Mapping Interface', level=2)
    phase4_text = """
    5. Dynamic Mapping Interface Creation (mapping_logic.py)
       • Create dedicated tab for each input file/sheet combination
       • Read and process input data with performance optimization
       • Handle column header detection and row offset
       • Strip whitespace and deduplicate column names
       • Parse enhanced mapping configuration
       • Build comprehensive mapping interface with:
         - Include/Exclude checkboxes
         - Output column display
         - Input column mapping dropdowns
         - Static value text inputs
         - Dynamic filter multiselects
         - Date formatting checkboxes
    """
    doc.add_paragraph(phase4_text)
    
    doc.add_heading('Phase 5: Data Processing & Output', level=2)
    phase5_text = """
    6. Final Output Processing (mapping_logic.py)
       • Apply user-selected filters to input data
       • Process column mappings with comprehensive logic
       • Handle static values, column mappings, and blank mappings
       • Apply date formatting when enabled
       • Validate all mappings for completeness
       • Combine all processed DataFrames into final result
       • Generate multiple export formats (Excel, TXT, Enhanced Mapping CSV)
    """
    doc.add_paragraph(phase5_text)
    
    # Key Processing Points
    doc.add_heading('4. Key Processing Points', level=1)
    
    doc.add_heading('4.1 Enhanced Mapping File Support', level=2)
    doc.add_paragraph(
        "The application provides full backward compatibility while enabling powerful new features "
        "through enhanced mapping files. Basic mapping files continue to work seamlessly, while "
        "enhanced files unlock advanced capabilities."
    )
    
    enhanced_features = [
        "Static Values: Assign fixed values to output columns",
        "Filter Values: Apply comma-separated filter criteria to input data", 
        "Date Format Flags: Control date formatting on a per-column basis",
        "Include Flags: Control column inclusion/exclusion in output"
    ]
    for feature in enhanced_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('4.2 Dynamic User Interface Generation', level=2)
    doc.add_paragraph(
        "The interface adapts dynamically to the uploaded data and configuration, providing "
        "contextual controls and real-time validation."
    )
    
    # User Journey Flow
    doc.add_heading('5. User Journey Flow', level=1)
    
    journey_steps = [
        "User Access: Navigate to Streamlit application with styled interface",
        "File Upload: Upload input files, output template, and optional mapping file",
        "Configuration: Select Excel sheets and configure mappings for each file/sheet",
        "Mapping Setup: Toggle column inclusion, map columns, assign static values, apply filters",
        "Output Generation: Click 'Generate Final Output' and receive validation feedback",
        "Export: Download consolidated Excel file, TXT file, and enhanced mapping configuration"
    ]
    
    for i, step in enumerate(journey_steps, 1):
        doc.add_paragraph(f"{i}. {step}", style='List Number')
    
    # Enhanced Features
    doc.add_heading('6. Enhanced Features', level=1)
    
    doc.add_heading('6.1 Advanced Mapping Capabilities', level=2)
    doc.add_paragraph(
        "The application supports sophisticated mapping configurations that go beyond "
        "simple column-to-column mapping, enabling complex data transformation workflows."
    )
    
    doc.add_heading('6.2 Configuration Management', level=2)
    doc.add_paragraph(
        "Enhanced mapping files provide comprehensive configuration export including all "
        "user settings, enabling team collaboration and process documentation."
    )
    
    # Example Enhanced Mapping Configuration
    doc.add_heading('Example Enhanced Mapping Configuration:', level=3)
    example_table = doc.add_table(rows=5, cols=8)
    example_table.style = 'Table Grid'
    
    # Header row
    header_cells = example_table.rows[0].cells
    headers = ['FileName', 'SheetName', 'OutputColumn', 'InputColumn', 'StaticValue', 'FilterValues', 'DateFormatFlag', 'IncludeFlag']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    data_rows = [
        ['data.xlsx', 'Sheet1', 'CustomerName', 'Full_Name', '', '', 'False', 'True'],
        ['data.xlsx', 'Sheet1', 'Status', '', 'Active', '', 'False', 'True'],
        ['data.xlsx', 'Sheet1', 'ProcessDate', 'Order_Date', '', '', 'True', 'True'],
        ['data.xlsx', 'Sheet1', 'Category', 'Product_Type', '', 'Premium,Standard', 'False', 'True']
    ]
    
    for row_idx, row_data in enumerate(data_rows, 1):
        cells = example_table.rows[row_idx].cells
        for col_idx, cell_data in enumerate(row_data):
            cells[col_idx].text = cell_data
    
    # Conclusion
    doc.add_heading('7. Conclusion', level=1)
    conclusion_text = (
        "This architecture provides a robust, scalable, and user-friendly solution for complex "
        "data mapping and transformation tasks while maintaining simplicity for basic use cases. "
        "The modular design enables easy maintenance and extension, while the enhanced mapping "
        "file format provides powerful configuration management capabilities for professional "
        "data processing workflows."
    )
    doc.add_paragraph(conclusion_text)
    
    # Footer
    doc.add_paragraph("\n" + "="*80)
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run("Document End\n\n").bold = True
    footer_para.add_run(
        "This document provides a comprehensive technical overview of the Advanced Column "
        "Mapping & Transformation Tool architecture. For additional technical details or "
        "specific implementation questions, refer to the source code documentation within each module."
    ).italic = True
    
    return doc

def main():
    """Main function to create and save the Word document."""
    print("Creating Word document...")
    
    try:
        # Create the document
        doc = create_word_document()
        
        # Save the document
        output_path = "Advanced_Column_Mapping_Tool_Architecture_Guide.docx"
        doc.save(output_path)
        
        print(f"✅ Word document created successfully: {output_path}")
        print(f"📄 Document location: {os.path.abspath(output_path)}")
        
        # Display file size
        file_size = os.path.getsize(output_path)
        print(f"📊 File size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
        
    except Exception as e:
        print(f"❌ Error creating Word document: {str(e)}")
        return False
    
    return True

if __name__ == "__main__":
    main()
